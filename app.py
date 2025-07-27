import streamlit as st
from io import BytesIO
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt

def html_to_docx(html_content):
    soup = BeautifulSoup(html_content, "html.parser")
    doc = Document()

    def add_table(table_tag):
        rows = table_tag.find_all("tr")
        if not rows:
            return
        cols = rows[0].find_all(["td", "th"])
        num_cols = len(cols)
        table = doc.add_table(rows=len(rows), cols=num_cols)
        table.style = "Table Grid"

        for i, row in enumerate(rows):
            cells = row.find_all(["td", "th"])
            for j, cell in enumerate(cells):
                # Set text inside each cell
                cell_text = cell.get_text(strip=True)
                table.cell(i, j).text = cell_text

    for elem in soup.body or soup:
        if elem.name in ["p", None]:  # paragraph or plain text
            text = elem.get_text(strip=True) if elem else str(elem)
            if text:
                doc.add_paragraph(text)
        elif elem.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
            level = int(elem.name[1])
            doc.add_heading(elem.get_text(strip=True), level=level-1)
        elif elem.name == "ul":
            for li in elem.find_all("li"):
                doc.add_paragraph(li.get_text(strip=True), style="ListBullet")
        elif elem.name == "ol":
            for li in elem.find_all("li"):
                doc.add_paragraph(li.get_text(strip=True), style="ListNumber")
        elif elem.name == "table":
            add_table(elem)

    docx_io = BytesIO()
    doc.save(docx_io)
    docx_io.seek(0)
    return docx_io

st.title("🌐 HTML to DOCX Converter with Table Support")

uploaded_file = st.file_uploader("Upload HTML file", type=["html", "htm"])

if uploaded_file:
    html_content = uploaded_file.read().decode("utf-8")

    if st.button("Convert to DOCX"):
        docx_file = html_to_docx(html_content)
        st.success("Conversion done!")
        st.download_button("📄 Download DOCX", docx_file, file_name="converted.docx")
